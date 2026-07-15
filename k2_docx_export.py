# -*- coding: utf-8 -*-
"""
Riktig DOCX-export med python-docx. Fungerar bara på Odoo.sh/egen server
eftersom det kräver ett Python-paket (external_dependencies i manifest).
"""

import io
import base64
import logging

from odoo import models, fields, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    _logger.warning(
        "python-docx är inte installerat. Kör 'pip install python-docx' "
        "för att aktivera DOCX-export av K2-årsredovisningar."
    )


class K2Arsredovisning(models.Model):
    _inherit = 'k2.arsredovisning'

    def action_export_docx(self):
        self.ensure_one()
        if Document is None:
            raise UserError(_(
                "python-docx är inte installerat på servern. "
                "Be er systemadministratör köra: pip install python-docx"
            ))

        doc = Document()

        # ── Stilar ──────────────────────────────────────────────────
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        def h1(text):
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x28, 0x55)

        def h2(text):
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x28, 0x55)

        # ── Försättssida ────────────────────────────────────────────
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.company_id.name)
        run.font.size = Pt(24)
        run.bold = True

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run('Org.nr %s' % (self.orgnr or '—'))

        h1_p = doc.add_paragraph()
        h1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h1_p.add_run('ÅRSREDOVISNING')
        run.font.size = Pt(18)
        run.bold = True

        period_p = doc.add_paragraph()
        period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_p.add_run('för räkenskapsåret %s – %s' % (self.date_from, self.date_to))

        doc.add_page_break()

        # ── Förvaltningsberättelse ──────────────────────────────────
        h1('Förvaltningsberättelse')
        h2('Verksamhetens art och inriktning')
        doc.add_paragraph(self.verksamhet or '[Ej ifyllt]')
        h2('Väsentliga händelser under räkenskapsåret')
        doc.add_paragraph(self.handelser or '[Ej ifyllt]')
        h2('Förväntad framtida utveckling')
        doc.add_paragraph(self.framtid or '[Ej ifyllt]')
        if self.miljo:
            h2('Miljö och personal')
            doc.add_paragraph(self.miljo)
        h2('Förslag till disposition av årets resultat')
        doc.add_paragraph(self.disposition or '[Ej ifyllt]')

        doc.add_page_break()

        # ── Resultaträkning ──────────────────────────────────────────
        h1('Resultaträkning')
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Post', str(self.date_to.year), 'Föreg. år'

        rr_rows = [
            ('Nettoomsättning', self.rr_netto, self.rr_p_netto),
            ('Övriga externa kostnader', self.rr_externa, self.rr_p_externa),
            ('Personalkostnader', self.rr_personal, self.rr_p_personal),
            ('Av- och nedskrivningar', self.rr_avskriv, self.rr_p_avskriv),
            ('Rörelseresultat', self.rr_rorelseresultat, self.rr_p_rorelseresultat),
            ('Resultat efter finansiella poster', self.rr_resultat_fin, self.rr_p_resultat_fin),
            ('ÅRETS RESULTAT', self.rr_arets_resultat, self.rr_p_arets_resultat),
        ]
        for label, n, f in rr_rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = '{:,.0f}'.format(n)
            row[2].text = '{:,.0f}'.format(f)

        doc.add_page_break()

        # ── Balansräkning ─────────────────────────────────────────────
        h1('Balansräkning')
        h2('TILLGÅNGAR')
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Light Grid Accent 1'
        hdr2 = table2.rows[0].cells
        hdr2[0].text, hdr2[1].text, hdr2[2].text = 'Post', str(self.date_to), 'Föreg. år'
        bt_rows = [
            ('Summa anläggningstillgångar', self.bt_anl_summa, self.bp_anl_summa),
            ('Summa omsättningstillgångar', self.bt_oms_summa, 0),
            ('SUMMA TILLGÅNGAR', self.bt_summa, self.bp_summa),
        ]
        for label, n, f in bt_rows:
            row = table2.add_row().cells
            row[0].text = label
            row[1].text = '{:,.0f}'.format(n)
            row[2].text = '{:,.0f}'.format(f) if f else '—'

        h2('EGET KAPITAL OCH SKULDER')
        table3 = doc.add_table(rows=1, cols=3)
        table3.style = 'Light Grid Accent 1'
        hdr3 = table3.rows[0].cells
        hdr3[0].text, hdr3[1].text, hdr3[2].text = 'Post', str(self.date_to), 'Föreg. år'
        be_rows = [
            ('Summa eget kapital', self.be_ek_summa, self.bp_ek_summa),
            ('Summa kortfristiga skulder', self.be_kort_summa, 0),
            ('SUMMA EK OCH SKULDER', self.be_summa, self.bp_ek_skuld_summa),
        ]
        for label, n, f in be_rows:
            row = table3.add_row().cells
            row[0].text = label
            row[1].text = '{:,.0f}'.format(n)
            row[2].text = '{:,.0f}'.format(f) if f else '—'

        doc.add_page_break()

        # ── Noter ───────────────────────────────────────────────────
        h1('Noter')
        h2('Not 1 – Redovisningsprinciper')
        doc.add_paragraph(self.not_principer or 'Upprättad enligt ÅRL och BFNAR 2016:10 (K2).')
        h2('Not 3 – Anställda och löner')
        doc.add_paragraph(self.not_personal or '[Ej ifyllt]')
        h2('Not 8 – Andelar i koncernföretag (dotterbolag)')
        doc.add_paragraph(self.not_dotterbolag or 'Bolaget äger inga andelar i dotterföretag.')
        h2('Not – Nyckeltalsdefinitioner')
        doc.add_paragraph(self.not_nyckeltal or '')

        doc.add_page_break()

        # ── Underskrifter ───────────────────────────────────────────
        h1('Underskrifter')
        for led in self.sign_styrelse_ids:
            doc.add_paragraph('%s, %s' % (led.name, dict(led._fields['befattning'].selection).get(led.befattning)))

        h1('Fastställelseintyg')
        doc.add_paragraph(
            'Undertecknad styrelseledamot i %s intygar dels att denna kopia av '
            'årsredovisningen stämmer överens med originalet, dels att resultaträkningen '
            'och balansräkningen har fastställts på årsstämma den %s. Årsstämman beslutade '
            'att godkänna styrelsens förslag till hur vinsten ska fördelas.'
            % (self.company_id.name, self.fast_stamma_datum or '[ej angivet]')
        )
        doc.add_paragraph('%s den %s' % (self.fast_ort, self.fast_datum or '[ej angivet]'))
        doc.add_paragraph(self.fast_styrelseledamot_id.name or '[Namnförtydligande]')

        # ── Spara som bilaga och returnera nedladdning ───────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = 'Arsredovisning_K2_%s_%s.docx' % (
            self.company_id.name.replace(' ', '_'), self.date_to.year
        )
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(buffer.read()),
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'self',
        }
